# ============================================================
#  SCHOLARA  —  AI Literature Review Workspace
#  Single-file Streamlit app  |  v3.0  |  Production-ready
# ============================================================
#
#  SETUP:
#    pip install streamlit pdfplumber python-docx requests google-generativeai
#    streamlit run app.py
#
#  REQUIRED SECRETS (.streamlit/secrets.toml):
#    GEMINI_API_KEY = "AIza..."
#    AIRTABLE_API_KEY = "key..."
#    AIRTABLE_BASE_ID = "app..."
# ============================================================

import io
import re
import time
import json
import requests
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import pdfplumber
import streamlit as st
from docx import Document
import google.generativeai as genai

# ── Airtable helpers ─────────────────────────────────────────
def airtable_create_record(table_name: str, fields: dict) -> Optional[str]:
    """Insert a row into Airtable. Returns the record ID."""
    try:
        api_key = st.secrets["AIRTABLE_API_KEY"]
        base_id = st.secrets["AIRTABLE_BASE_ID"]
    except Exception:
        st.warning("Airtable not configured – reviews saved in session only.")
        return None

    url = f"https://api.airtable.com/v0/{base_id}/{table_name}"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    data = {"records": [{"fields": fields}]}
    resp = requests.post(url, json=data, headers=headers)
    if resp.status_code == 200:
        return resp.json()["records"][0]["id"]
    else:
        st.warning(f"Airtable error: {resp.text}")
        return None


def airtable_get_records(table_name: str, user_id: str = None) -> list:
    """Fetch all records, optionally filtered by user_id."""
    try:
        api_key = st.secrets["AIRTABLE_API_KEY"]
        base_id = st.secrets["AIRTABLE_BASE_ID"]
    except Exception:
        return []

    url = f"https://api.airtable.com/v0/{base_id}/{table_name}"
    headers = {"Authorization": f"Bearer {api_key}"}
    params = {}
    if user_id:
        params["filterByFormula"] = f"{{user_id}}='{user_id}'"
    resp = requests.get(url, headers=headers, params=params)
    if resp.status_code == 200:
        return resp.json().get("records", [])
    else:
        st.warning(f"Airtable read error: {resp.text}")
        return []


def get_user_id() -> str:
    if "user_id" not in st.session_state:
        st.session_state.user_id = "guest"
    return st.session_state.user_id


# ── 1. PAGE CONFIG ───────────────────────────────────────────
st.set_page_config(
    page_title="Scholara — AI Literature Review",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── 2. CONSTANTS ─────────────────────────────────────────────
APP_NAME    = "Scholara"
APP_VERSION = "3.0.0"
APP_TAGLINE = "AI Literature Review Workspace"

MAX_FILES      = 50
MAX_TOTAL_MB   = 120
MAX_TEXT_CHARS = 160_000
PREVIEW_CHARS  = 320

CITATION_STYLES = ["APA 7", "IEEE", "MLA 9", "Chicago 17", "Vancouver"]
REVIEW_LENGTHS  = ["Short (~700 words)", "Medium (~1,200 words)", "Detailed (~1,800 words)"]
TONES           = ["Academic", "Concise", "Critical", "Neutral"]
SORT_OPTIONS    = ["Newest first", "Oldest first", "Title A-Z", "Title Z-A", "Most papers"]

# ── 3. SESSION STATE ─────────────────────────────────────────
def _init_state() -> None:
    defaults: Dict[str, Any] = {
        "page":        "Home",
        "reviews":     [],
        "cur_review":  None,
        "cur_meta":    {},
        "search":      "",
        "sort_by":     "Newest first",
        "show_tips":   True,
        "last_error":  "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init_state()

# ── 4. HELPERS ───────────────────────────────────────────────
def now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")

def date_slug() -> str:
    return datetime.now().strftime("%Y%m%d")

def slugify(text: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9\s_-]", "", text).strip().lower()
    s = re.sub(r"[\s_-]+", "_", s)
    return (s[:60] or "review")

def file_kb(size_bytes: int) -> float:
    return round(size_bytes / 1024, 1)

def total_mb(files: List[Any]) -> float:
    return sum(f.size for f in files) / (1024 * 1024)

def word_count(text: str) -> int:
    return len(text.split())

def nav_to(page: str) -> None:
    st.session_state.page = page
    st.rerun()

# ── 5. DATA / GENERATION HELPERS ────────────────────────────
def validate_uploads(files: List[Any]) -> Dict[str, Any]:
    if not files:
        return {"ok": False, "msg": "Upload at least one PDF."}
    if len(files) > MAX_FILES:
        return {"ok": False, "msg": f"Maximum {MAX_FILES} PDFs per review."}
    mb = total_mb(files)
    if mb > MAX_TOTAL_MB:
        return {"ok": False, "msg": f"Total {mb:.1f} MB exceeds {MAX_TOTAL_MB} MB limit."}
    names = [f.name for f in files]
    dupes = sorted({n for n in names if names.count(n) > 1})
    if dupes:
        return {"ok": False, "msg": f"Duplicate filenames: {', '.join(dupes)}"}
    return {"ok": True, "msg": f"{len(files)} file(s) ready  •  {mb:.1f} MB total", "mb": mb}


@st.cache_data(show_spinner=False)
def _extract_pdf(name: str, data: bytes) -> str:
    blocks: List[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                blocks.append(txt.strip())
    return "\n".join(blocks)


def extract_all(files: List[Any]) -> Tuple[str, List[str]]:
    parts: List[str] = []
    warnings: List[str] = []
    for f in files:
        try:
            raw = f.read()
            if not raw:
                warnings.append(f"{f.name}: empty file.")
                continue
            txt = _extract_pdf(f.name, raw)
            if txt.strip():
                parts.append(txt)
            else:
                warnings.append(f"{f.name}: no extractable text (scanned/protected?).")
        except Exception as ex:
            warnings.append(f"{f.name}: parse error — {ex}")
    combined = "\n\n".join(parts).strip()
    if len(combined) > MAX_TEXT_CHARS:
        combined = combined[:MAX_TEXT_CHARS] + "\n\n[TRUNCATED]"
        warnings.append(f"Input capped at {MAX_TEXT_CHARS:,} chars for stability.")
    return combined, warnings


def build_prompt(
    text: str, topic: str, citation: str, length: str, tone: str,
    inc_lim: bool, inc_gaps: bool, inc_table: bool, inc_impl: bool,
) -> str:
    secs = [
        "1. Introduction",
        "2. Thematic Synthesis",
        "3. Methodological Trends",
        "4. Contradictions & Debates",
    ]
    if inc_lim:   secs.append("5. Limitations of Existing Literature")
    if inc_gaps:  secs.append("6. Research Gaps & Future Directions")
    if inc_table: secs.append("7. Methods Comparison Table (plain text)")
    if inc_impl:  secs.append("8. Practical Implications")
    secs.append("9. Conclusion")
    secs.append(f"10. Inline citations ({citation} style)")

    topic_line = topic.strip() or "Infer the unifying topic from the uploaded papers."

    return (
        "You are an expert academic research assistant.\n\n"
        "TASK: Write a structured literature review using ONLY the source text below.\n\n"
        f"PARAMETERS:\n"
        f"- Topic: {topic_line}\n"
        f"- Citation style: {citation}\n"
        f"- Length: {length}\n"
        f"- Tone: {tone}\n\n"
        "RULES:\n"
        "- Never fabricate studies, authors, or data.\n"
        "- If evidence for a claim is absent, write: "
        "'insufficient evidence in uploaded papers.'\n"
        "- Use clear headings for each section.\n\n"
        "OUTPUT SECTIONS:\n"
        + "\n".join(secs)
        + "\n\nSOURCE TEXT:\n"
        + text
    )


def _call_api(prompt: str) -> Optional[str]:
    try:
        api_key = st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        api_key = ""

    if not api_key:
        time.sleep(0.8)
        return (
            "## Introduction\n"
            "This literature review synthesises the uploaded papers, identifying key "
            "themes, methodological patterns, and research gaps.\n\n"
            "## Thematic Synthesis\n"
            "Three recurring themes emerge: performance benchmarking, deployment "
            "constraints, and evaluation inconsistency across datasets.\n\n"
            "## Methodological Trends\n"
            "Most studies rely on controlled benchmark experiments with limited "
            "external validity or longitudinal follow-up.\n\n"
            "## Contradictions & Debates\n"
            "Findings diverge across evaluation metrics and dataset choices, reducing "
            "cross-study comparability.\n\n"
            "## Research Gaps & Future Directions\n"
            "Standardised reporting protocols and real-world longitudinal validation "
            "are underrepresented and represent clear opportunities.\n\n"
            "## Conclusion\n"
            "The field shows strong momentum. Translating benchmark results into "
            "reliable real-world performance remains the primary open challenge.\n\n"
            "_This is demo output. Add GEMINI\\_API\\_KEY to `.streamlit/secrets.toml` "
            "for full generation._"
        )

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-3.5-flash-lite')
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.session_state.last_error = str(e)
        return None


def generate_review(
    text: str, topic: str, citation: str, length: str, tone: str,
    inc_lim: bool, inc_gaps: bool, inc_table: bool, inc_impl: bool,
) -> Optional[str]:
    try:
        prompt = build_prompt(
            text=text, topic=topic, citation=citation,
            length=length, tone=tone,
            inc_lim=inc_lim, inc_gaps=inc_gaps,
            inc_table=inc_table, inc_impl=inc_impl,
        )
        return _call_api(prompt)
    except Exception as ex:
        st.session_state.last_error = str(ex)
        return None


def export_txt(text: str) -> bytes:
    return text.encode("utf-8")


def export_docx(text: str, title: str = "Literature Review") -> bytes:
    doc = Document()
    doc.add_heading(title, 1)
    for block in text.split("\n\n"):
        b = block.strip()
        if not b:
            continue
        if b.startswith("### "):
            doc.add_heading(b[4:].strip(), 3)
        elif b.startswith("## "):
            doc.add_heading(b[3:].strip(), 2)
        elif b.startswith("# "):
            doc.add_heading(b[2:].strip(), 1)
        else:
            doc.add_paragraph(b)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def meta_chips_html(meta: Dict) -> str:
    bits = []
    for k in ["citation", "tone", "length"]:
        v = meta.get(k)
        if v:
            bits.append(f"<span class='chip'>{v}</span>")
    if meta.get("inc_gaps"):  bits.append("<span class='chip chip-green'>Gaps</span>")
    if meta.get("inc_lim"):   bits.append("<span class='chip'>Limitations</span>")
    if meta.get("inc_table"): bits.append("<span class='chip chip-amber'>Methods Table</span>")
    if meta.get("inc_impl"):  bits.append("<span class='chip'>Implications</span>")
    return "".join(bits)


def save_review(title: str, text: str, papers: List[str], meta: Dict) -> None:
    entry = {
        "title": title.strip() or f"Review {now_str()}",
        "date": now_str(),
        "text": text,
        "papers": papers,
        "meta": meta,
    }
    st.session_state.reviews.append(entry)
    st.session_state.cur_review = text
    st.session_state.cur_meta   = meta

    fields = {
        "user_id": get_user_id(),
        "title": entry["title"],
        "review_text": text,
        "papers": json.dumps(papers),
        "meta": json.dumps(meta),
        "created_at": now_str()
    }
    record_id = airtable_create_record("reviews", fields)
    if record_id:
        st.toast("✅ Saved to cloud (Airtable)!")
    else:
        st.toast("⚠️ Saved locally only – check secrets")


# ── 6. CSS ────────────────────────────────────────────────────
def _inject_css() -> None:
    st.markdown(
        """
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display&family=Inter:wght@400;500;600;700;800&display=swap');

:root {
  --ink:       #0D1B2A;
  --navy:      #1B2E4B;
  --navy-mid:  #243B55;
  --indigo:    #4361EE;
  --indigo-lt: #7B93FF;
  --indigo-dim:#EEF1FF;
  --amber:     #F4A261;
  --amber-lt:  #FFF3E8;
  --slate-50:  #F8FAFC;
  --slate-100: #F1F5F9;
  --slate-200: #E2E8F0;
  --slate-400: #94A3B8;
  --slate-600: #475569;
  --white:     #FFFFFF;
  --success:   #16A34A;
  --sh-sm: 0 1px 3px rgba(13,27,42,.08),0 1px 2px rgba(13,27,42,.06);
  --sh-md: 0 4px 12px rgba(13,27,42,.10),0 2px 6px rgba(13,27,42,.06);
  --sh-lg: 0 12px 32px rgba(13,27,42,.12),0 4px 12px rgba(13,27,42,.06);
  --r-xl: 20px; --r-lg: 16px; --r-md: 10px; --r-sm: 6px;
}

html, body { font-family:'Inter',system-ui,sans-serif !important; background:#F8FAFC !important; }
.stApp     { background:#F8FAFC !important; font-family:'Inter',system-ui,sans-serif !important; }
div.block-container { max-width:1160px !important; padding-top:1.2rem !important; }

section[data-testid="stSidebar"] {
  background: var(--navy) !important;
  border-right: 1px solid var(--navy-mid);
}
section[data-testid="stSidebar"] * { color:#CBD5E1 !important; }
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] strong { color:#FFFFFF !important; }
section[data-testid="stSidebar"] .stRadio label { color:#94A3B8 !important; font-size:0.87rem; }

::-webkit-scrollbar { width:5px; }
::-webkit-scrollbar-thumb { background:var(--slate-400); border-radius:3px; }

.skip-link {
  position:absolute; top:-48px; left:8px;
  background:var(--indigo); color:#fff !important;
  padding:8px 14px; border-radius:var(--r-md);
  font-size:0.84rem; z-index:9999; text-decoration:none;
}
.skip-link:focus { top:8px; }

.hero-wrap {
  background: linear-gradient(135deg, var(--navy) 0%, var(--navy-mid) 60%, #1B3A6B 100%);
  border-radius: var(--r-xl); padding: 52px 48px 44px;
  margin-bottom: 24px; position: relative; overflow: hidden;
  box-shadow: var(--sh-lg);
}
.hero-wrap::before {
  content:''; position:absolute; top:-60px; right:-60px;
  width:320px; height:320px;
  background: radial-gradient(circle,rgba(67,97,238,.4) 0%,transparent 70%);
  border-radius:50%; pointer-events:none;
}
.hero-eyebrow {
  display:inline-block; font-size:0.70rem; font-weight:700;
  letter-spacing:0.12em; text-transform:uppercase;
  color: var(--amber) !important; background: rgba(244,162,97,.15);
  border:1px solid rgba(244,162,97,.3); border-radius:999px;
  padding:4px 14px; margin-bottom:16px;
}
.hero-title {
  font-family:'DM Serif Display',Georgia,serif !important;
  font-size:2.6rem !important; line-height:1.18 !important;
  color:#FFFFFF !important; margin:0 0 12px !important;
}
.hero-sub {
  font-size:1.02rem; color:#94A3B8 !important;
  line-height:1.6; max-width:520px; margin:0;
}

.trust-strip { display:flex; flex-wrap:wrap; gap:8px; margin-bottom:24px; }
.trust-item {
  display:flex; align-items:center; gap:6px;
  background:var(--white); border:1px solid var(--slate-200);
  border-radius:999px; padding:5px 13px;
  font-size:0.80rem; color:var(--slate-600) !important; font-weight:500;
  box-shadow:var(--sh-sm);
}

.kpi-grid { display:grid; grid-template-columns:repeat(3,1fr); gap:14px; margin-bottom:24px; }
.kpi-card {
  background:var(--white); border:1px solid var(--slate-200);
  border-radius:var(--r-lg); padding:18px 20px; box-shadow:var(--sh-sm);
  position:relative; overflow:hidden;
}
.kpi-card::before {
  content:''; position:absolute; top:0; left:0; right:0; height:3px;
  background:linear-gradient(90deg,var(--indigo),var(--indigo-lt));
}
.kpi-label { font-size:0.76rem; font-weight:600; letter-spacing:0.06em;
             text-transform:uppercase; color:var(--slate-400) !important; margin-bottom:6px; }
.kpi-value { font-family:'DM Serif Display',Georgia,serif;
             font-size:1.6rem; color:var(--ink) !important; line-height:1; }
.kpi-sub   { font-size:0.76rem; color:var(--slate-400) !important; margin-top:2px; }

.s-heading {
  font-family:'DM Serif Display',Georgia,serif;
  font-size:1.45rem; color:var(--ink) !important;
  margin:0 0 4px; line-height:1.25;
}
.s-caption { font-size:0.86rem; color:var(--slate-400) !important; margin-bottom:18px; }

.s-card {
  background:var(--white); border:1px solid var(--slate-200);
  border-radius:var(--r-lg); padding:20px 22px;
  box-shadow:var(--sh-sm); margin-bottom:12px;
}
.s-card-accent { border-left:4px solid var(--indigo); padding-left:18px; }
.s-card-dark   { background:var(--navy); border-color:var(--navy-mid); }
.s-card-amber  { background:var(--amber-lt); border-color:#FCD9B6;
                 border-left:4px solid var(--amber); padding-left:18px; }
.s-card-dash   { border:1px dashed var(--slate-200); background:var(--slate-50);
                 text-align:center; padding:40px 24px; }

.step-bar { display:grid; grid-template-columns:repeat(4,1fr); gap:6px; margin-bottom:20px; }
.step-item {
  background:var(--white); border:1px solid var(--slate-200);
  border-radius:var(--r-md); padding:9px 10px;
  text-align:center; font-size:0.78rem; font-weight:600;
  color:var(--slate-400) !important;
}
.step-item.active { background:var(--indigo-dim); border-color:var(--indigo); color:var(--indigo) !important; }
.step-item.done   { background:#ECFDF5; border-color:var(--success); color:var(--success) !important; }

.chip {
  display:inline-block; background:var(--indigo-dim); color:var(--indigo) !important;
  border:1px solid #C7D0FF; border-radius:999px;
  padding:3px 10px; font-size:0.74rem; font-weight:600;
  margin-right:4px; margin-bottom:4px;
}
.chip-amber { background:var(--amber-lt); color:#92400E !important; border-color:#FCD9B6; }
.chip-green { background:#ECFDF5; color:var(--success) !important; border-color:#BBF7D0; }

.rev-card { background:var(--white); border:1px solid var(--slate-200);
            border-radius:var(--r-lg); padding:16px 20px;
            margin-bottom:10px; box-shadow:var(--sh-sm); }
.rev-title { font-family:'DM Serif Display',Georgia,serif;
             font-size:1.10rem; color:var(--ink) !important; margin-bottom:3px; }
.rev-meta  { font-size:0.78rem; color:var(--slate-400) !important; margin-bottom:8px; }
.rev-preview { font-size:0.85rem; color:var(--slate-600) !important; line-height:1.55;
               border-left:3px solid var(--slate-200); padding-left:10px; margin:8px 0 12px; }

.price-card {
  background:var(--white); border:1px solid var(--slate-200);
  border-radius:var(--r-xl); padding:26px 22px; box-shadow:var(--sh-sm);
}
.price-card.featured { border-color:var(--indigo);
  box-shadow:0 0 0 3px rgba(67,97,238,.14),var(--sh-md); }
.price-name { font-size:0.74rem; font-weight:700; letter-spacing:0.10em;
              text-transform:uppercase; color:var(--indigo) !important; margin-bottom:5px; }
.price-val  { font-family:'DM Serif Display',Georgia,serif;
              font-size:2.2rem; color:var(--ink) !important; line-height:1; }
.price-per  { font-size:0.82rem; color:var(--slate-400) !important; margin-bottom:16px; }
.price-row  { display:flex; align-items:center; gap:8px; padding:5px 0;
              border-bottom:1px solid var(--slate-100); font-size:0.83rem; }

.blog-card { background:var(--white); border:1px solid var(--slate-200);
             border-radius:var(--r-lg); padding:18px 20px; margin-bottom:10px;
             display:flex; align-items:flex-start; gap:14px; box-shadow:var(--sh-sm); }
.blog-num  { font-family:'DM Serif Display',Georgia,serif;
             font-size:1.4rem; color:var(--slate-200) !important; min-width:32px; }
.blog-title{ font-weight:600; color:var(--ink) !important; font-size:0.93rem; margin-bottom:3px; }
.blog-meta { font-size:0.76rem; color:var(--slate-400) !important; }

.tip-box {
  background:var(--indigo-dim); border:1px solid #C7D0FF;
  border-left:4px solid var(--indigo); border-radius:var(--r-md);
  padding:9px 13px; font-size:0.83rem; color:var(--navy-mid) !important;
  margin-bottom:14px;
}

.s-div { height:1px; background:var(--slate-200); margin:22px 0; }

.page-footer {
  text-align:center; font-size:0.76rem; color:var(--slate-400) !important;
  margin-top:44px; padding:16px 0 6px; border-top:1px solid var(--slate-200);
}

.sb-brand { padding:14px 0 10px; border-bottom:1px solid rgba(255,255,255,.08); margin-bottom:10px; }
.sb-name  { font-family:'DM Serif Display',Georgia,serif; font-size:1.25rem; color:#FFFFFF !important; }
.sb-tag   { font-size:0.74rem; color:#475569 !important; margin-top:2px; }
.sb-group { font-size:0.67rem; font-weight:700; letter-spacing:0.10em;
            text-transform:uppercase; color:#475569 !important; padding:12px 0 4px; }
.sb-stat  { display:flex; justify-content:space-between; font-size:0.78rem;
            padding:4px 0; border-bottom:1px solid rgba(255,255,255,.06); color:#94A3B8 !important; }
.sb-val   { font-weight:700; color:#FFFFFF !important; }

div[data-testid="stProgress"] > div > div > div {
  background:linear-gradient(90deg,var(--indigo),var(--indigo-lt)) !important;
  border-radius:999px !important;
}
.stButton > button {
  border-radius:var(--r-md) !important;
  font-family:'Inter',sans-serif !important;
  font-size:0.87rem !important; font-weight:600 !important;
}
.stButton > button[kind="primary"] {
  background:linear-gradient(135deg,var(--indigo),#3A56D8) !important;
  border:none !important; color:#fff !important;
}
.stTextInput > div > div > input,
.stSelectbox > div > div,
.stTextArea > div > div > textarea {
  border-radius:var(--r-md) !important;
  border:1px solid var(--slate-200) !important;
}
button:focus, a:focus, input:focus { outline:3px solid var(--indigo) !important; outline-offset:2px !important; }

@media(max-width:768px) {
  .hero-title { font-size:1.8rem !important; }
  .hero-wrap  { padding:34px 24px 28px; }
  .kpi-grid   { grid-template-columns:1fr; }
  .step-bar   { grid-template-columns:repeat(2,1fr); }
}
</style>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(
        '<a class="skip-link" href="#main-content" tabindex="0">Skip to main content</a>'
        '<span id="main-content"></span>',
        unsafe_allow_html=True,
    )

_inject_css()

# ── 7. SHARED UI HELPERS ─────────────────────────────────────
def tip(msg: str) -> None:
    if st.session_state.show_tips:
        st.markdown(f"<div class='tip-box'>💡 {msg}</div>", unsafe_allow_html=True)

def divider() -> None:
    st.markdown("<div class='s-div'></div>", unsafe_allow_html=True)

def page_head(title: str, caption: str = "") -> None:
    st.markdown(f"<h2 class='s-heading'>{title}</h2>", unsafe_allow_html=True)
    if caption:
        st.markdown(f"<p class='s-caption'>{caption}</p>", unsafe_allow_html=True)

def kpi_row(items: List[Tuple[str, str, str]]) -> None:
    cards = "".join(
        f"<div class='kpi-card'>"
        f"<div class='kpi-label'>{lbl}</div>"
        f"<div class='kpi-value'>{val}</div>"
        f"<div class='kpi-sub'>{sub}</div>"
        f"</div>"
        for lbl, val, sub in items
    )
    st.markdown(f"<div class='kpi-grid'>{cards}</div>", unsafe_allow_html=True)

def step_bar(current: int) -> None:
    labels = ["Upload", "Configure", "Generate", "Export"]
    items = ""
    for i, lbl in enumerate(labels, 1):
        if i < current:
            cls = "step-item done"
            ico = "✓"
        elif i == current:
            cls = "step-item active"
            ico = str(i)
        else:
            cls = "step-item"
            ico = str(i)
        items += f"<div class='{cls}'>{ico} {lbl}</div>"
    st.markdown(f"<div class='step-bar'>{items}</div>", unsafe_allow_html=True)
    st.progress(min((current - 1) / 4, 1.0))

# ── 8. SIDEBAR ───────────────────────────────────────────────
def _sidebar() -> None:
    st.sidebar.markdown(
        "<div class='sb-brand'>"
        "<div class='sb-name'>📚 " + APP_NAME + "</div>"
        "<div class='sb-tag'>" + APP_TAGLINE + " · v" + APP_VERSION + "</div>"
        "</div>",
        unsafe_allow_html=True,
    )

    st.sidebar.markdown("<div class='sb-group'>Workspace</div>", unsafe_allow_html=True)
    st.sidebar.markdown("", unsafe_allow_html=True)
    st.sidebar.markdown("<div class='sb-group'>Product</div>", unsafe_allow_html=True)
    st.sidebar.markdown("", unsafe_allow_html=True)
    st.sidebar.markdown("<div class='sb-group'>Company</div>", unsafe_allow_html=True)

    all_pages = ["Home", "Demo Wizard", "Dashboard",
                 "How It Works", "Features", "Use Cases", "Pricing",
                 "About", "Blog"]

    current = st.session_state.page
    try:
        idx = all_pages.index(current)
    except ValueError:
        idx = 0

    selected = st.sidebar.radio(
        "Navigation",
        all_pages,
        index=idx,
        label_visibility="collapsed",
        key="nav_flat",
    )

    if selected != current:
        st.session_state.page = selected
        st.rerun()

    st.sidebar.markdown("<div class='sb-group'>Quick actions</div>", unsafe_allow_html=True)
    qa1, qa2 = st.sidebar.columns(2)
    with qa1:
        if st.button("New", key="sb_new", use_container_width=True):
            st.session_state.cur_review = None
            st.session_state.cur_meta   = {}
            st.session_state.page       = "Demo Wizard"
            st.rerun()
    with qa2:
        if st.button("Last", key="sb_last", use_container_width=True):
            if st.session_state.reviews:
                last = st.session_state.reviews[-1]
                st.session_state.cur_review = last["text"]
                st.session_state.cur_meta   = last.get("meta", {})
                st.session_state.page       = "Demo Wizard"
                st.rerun()
            else:
                st.sidebar.info("No reviews yet.")

    with st.sidebar.expander("User", expanded=False):
        current_user = st.session_state.get("user_id", "guest")
        new_id = st.text_input(
            "Your ID (e.g., soban_nust)",
            value=current_user,
            key="user_id_input"
        )
        if new_id != current_user:
            st.session_state.user_id = new_id
            st.rerun()
        st.caption("Reviews are stored under this ID.")

    st.sidebar.markdown("<div class='sb-group'>Session</div>", unsafe_allow_html=True)
    n_rev    = len(st.session_state.reviews)
    n_papers = sum(len(r.get("papers", [])) for r in st.session_state.reviews)
    last_run = st.session_state.reviews[-1]["date"] if st.session_state.reviews else "—"
    st.sidebar.markdown(
        "<div style='padding:2px 0;'>"
        "<div class='sb-stat'><span>Reviews</span><span class='sb-val'>" + str(n_rev) + "</span></div>"
        "<div class='sb-stat'><span>Papers</span><span class='sb-val'>" + str(n_papers) + "</span></div>"
        "<div class='sb-stat'><span>Last run</span><span class='sb-val' style='font-size:.72rem;'>"
        + last_run + "</span></div></div>",
        unsafe_allow_html=True,
    )

    with st.sidebar.expander("Preferences", expanded=False):
        st.session_state.show_tips = st.checkbox(
            "Show UI tips", value=st.session_state.show_tips)
        st.caption("Dark mode controls coming in v3.1.")

    st.sidebar.markdown("---")
    st.sidebar.caption("Cloud-backed (Airtable)")
    st.sidebar.caption("© 2026 Scholara")

# ── 9. HOME ──────────────────────────────────────────────────
# ── 9. HOME ──────────────────────────────────────────────────
def home_page() -> None:
    # callbacks that run before the next rerun
    def go_demo():
        st.session_state.page = "Demo Wizard"

    def go_dash():
        st.session_state.page = "Dashboard"

    def go_how():
        st.session_state.page = "How It Works"

    # Hero
    st.markdown(
        "<div class='hero-wrap'>"
        "<div class='hero-eyebrow'>✦ AI-Powered Research Tool</div>"
        "<h1 class='hero-title'>Your literature review,<br>written in minutes.</h1>"
        "<p class='hero-sub'>Upload your research papers. Scholara reads them, "
        "synthesises the themes, identifies gaps, and produces a structured, "
        "citation-ready draft — so you can focus on the thinking that matters.</p>"
        "</div>",
        unsafe_allow_html=True,
    )
    st.markdown(
        "<div class='trust-strip'>"
        "<span class='trust-item'>☁️ Cloud-backed (Airtable)</span>"
        "<span class='trust-item'>📚 Source-grounded synthesis</span>"
        "<span class='trust-item'>📝 Export to TXT &amp; DOCX</span>"
        "<span class='trust-item'>🎓 APA · IEEE · MLA · Chicago · Vancouver</span>"
        "</div>",
        unsafe_allow_html=True,
    )
    kpi_row([
        ("Time saved / review", "30–40 hrs", "per literature review"),
        ("Upload volume",       "50 PDFs",   "per review session"),
        ("Export formats",      "TXT + DOCX","open directly in Word"),
    ])

    c1, c2, c3 = st.columns([1.1, 1, 1])
    with c1:
        st.button("🚀 Start Demo Wizard", type="primary", use_container_width=True,
                  key="home_demo", on_click=go_demo)
    with c2:
        st.button("📊 Open Dashboard", use_container_width=True,
                  key="home_dash", on_click=go_dash)
    with c3:
        st.button("📖 How It Works", use_container_width=True,
                  key="home_how", on_click=go_how)

    divider()
    page_head("Built for researchers who are tired of formatting")
    a, b = st.columns(2)
    with a:
        st.markdown(
            "<div class='s-card s-card-accent'><strong>Guided 4-step wizard</strong>"
            "<p style='font-size:.86rem;color:var(--slate-600);margin-top:6px;'>"
            "Upload → Configure → Generate → Export. Each step confirmed before the next.</p></div>"
            "<div class='s-card s-card-accent'><strong>Configurable depth &amp; tone</strong>"
            "<p style='font-size:.86rem;color:var(--slate-600);margin-top:6px;'>"
            "Short to Detailed. Academic, Concise, Critical, or Neutral.</p></div>",
            unsafe_allow_html=True,
        )
    with b:
        st.markdown(
            "<div class='s-card s-card-accent'><strong>Searchable dashboard</strong>"
            "<p style='font-size:.86rem;color:var(--slate-600);margin-top:6px;'>"
            "All reviews in one place — searchable, sortable, re-downloadable.</p></div>"
            "<div class='s-card s-card-accent'><strong>Source-grounded output</strong>"
            "<p style='font-size:.86rem;color:var(--slate-600);margin-top:6px;'>"
            "AI synthesises only from your uploads. Missing evidence is flagged, not invented.</p></div>",
            unsafe_allow_html=True,
        )

    divider()
    page_head("Trusted by researchers at NUST, IBA, and LUMS")
    t1, t2, t3 = st.columns(3)
    quotes = [
        ("Scholara saved me two weeks on my thesis Chapter 2.", "Ayesha R.", "PhD Candidate, NUST"),
        ("I use it weekly for grant backgrounds. Genuine time-saver.", "Dr. Usman K.", "Asst. Professor, IBA"),
        ("The clearest AI review tool I have used — output actually makes sense.", "Hamza M.", "Research Assistant, LUMS"),
    ]
    for col, (q, name, role) in zip([t1, t2, t3], quotes):
        with col:
            st.markdown(
                "<div class='s-card' style='height:100%;'>"
                "<p style='font-size:.87rem;color:var(--slate-600);font-style:italic;"
                "line-height:1.6;margin-bottom:10px;'>\"" + q + "\"</p>"
                "<strong style='font-size:.83rem;'>" + name + "</strong><br/>"
                "<span style='font-size:.76rem;color:var(--slate-400);'>" + role + "</span>"
                "</div>",
                unsafe_allow_html=True,
            )
    st.markdown("<div class='page-footer'>© 2026 Scholara · v" + APP_VERSION + "</div>",
                unsafe_allow_html=True)
# ── 10. DEMO WIZARD ──────────────────────────────────────────
def demo_wizard_page() -> None:
    page_head("Demo Wizard", "Guided 4-step workflow — each step confirmed before the next.")

    step = 1
    if st.session_state.cur_review:
        step = 4
    elif st.session_state.get("_wiz_configured"):
        step = 3

    step_bar(step)
    divider()

    # Step 1: Upload
    page_head("Step 1 — Upload your papers",
              "Add up to 50 PDFs. Text-based PDFs extract best.")
    tip("Run OCR first on scanned image PDFs for reliable text extraction.")

    uploaded = st.file_uploader(
        "Choose PDF files",
        type=["pdf"],
        accept_multiple_files=True,
        key="wizard_uploader",
        label_visibility="collapsed",
    )

    files_ok = False
    if uploaded:
        v = validate_uploads(uploaded)
        if v["ok"]:
            st.success(v["msg"])
            files_ok = True
            with st.expander("View uploaded files", expanded=False):
                for f in uploaded:
                    st.write(f"• {f.name}  ({file_kb(f.size)} KB)")
        else:
            st.error(v["msg"])
            return
    else:
        st.markdown(
            "<div class='s-card s-card-dash'>"
            "<div style='font-size:2rem;margin-bottom:8px;'>📄</div>"
            "<strong style='color:var(--slate-600);'>Drop PDFs here or use the picker above</strong>"
            "</div>",
            unsafe_allow_html=True,
        )
        return

    divider()

    # Step 2: Configure
    page_head("Step 2 — Configure your review")

    cl, cr = st.columns([1.05, 1], gap="large")
    with cl:
        topic   = st.text_input("Research topic / question",
                                placeholder="e.g., Federated Learning for Healthcare")
        citation = st.selectbox("Citation style", CITATION_STYLES)
        length   = st.selectbox("Review depth",   REVIEW_LENGTHS, index=1)
        tone     = st.selectbox("Writing tone",   TONES)
    with cr:
        st.markdown("<p style='font-size:.76rem;font-weight:700;letter-spacing:.06em;"
                    "text-transform:uppercase;color:var(--slate-400);margin-bottom:6px;'>"
                    "Optional sections</p>", unsafe_allow_html=True)
        inc_lim   = st.checkbox("Limitations of existing literature", value=True)
        inc_gaps  = st.checkbox("Research gaps & future directions",  value=True)
        inc_table = st.checkbox("Methods comparison table",           value=False)
        inc_impl  = st.checkbox("Practical implications",             value=False)

        preview = (
            f"<span class='chip'>{citation}</span>"
            f"<span class='chip'>{tone}</span>"
            + ("<span class='chip chip-green'>Gaps</span>" if inc_gaps else "")
            + ("<span class='chip'>Limitations</span>"    if inc_lim  else "")
            + ("<span class='chip chip-amber'>Methods Table</span>" if inc_table else "")
            + ("<span class='chip'>Implications</span>"   if inc_impl else "")
        )
        st.markdown(
            "<div class='s-card' style='margin-top:10px;padding:12px 14px;'>"
            "<p style='font-size:.73rem;font-weight:700;letter-spacing:.06em;"
            "text-transform:uppercase;color:var(--slate-400);margin-bottom:6px;'>"
            "Your review will include</p>" + preview + "</div>",
            unsafe_allow_html=True,
        )

    divider()

    # Step 3: Generate
    page_head("Step 3 — Generate")
    gc, nc = st.columns([1, 1.5], gap="large")
    with gc:
        gen_clicked = st.button("🧠 Generate Literature Review",
                                type="primary", use_container_width=True)
    with nc:
        st.markdown("<div class='tip-box' style='margin:0;'>"
                    "⏱️ Powered by Gemini — output grounded in your papers only.</div>",
                    unsafe_allow_html=True)

    if gen_clicked:
        all_text, warnings = "", []

        with st.status("Working on your review…", expanded=True) as status:
            status.write("Phase 1/3 · Validating files…")
            time.sleep(0.2)

            status.write("Phase 2/3 · Extracting text from PDFs…")
            all_text, warnings = extract_all(uploaded)

            if warnings:
                for w in warnings[:5]:
                    status.write(f"  ⚠ {w}")

            if not all_text.strip():
                status.update(label="Extraction failed.", state="error")
                st.error("No extractable text found. Try OCR-processed PDFs.")
                return

            status.write(f"  ✓ {len(all_text):,} characters from {len(uploaded)} file(s).")
            status.write("Phase 3/3 · Generating with Gemini…")

            result = generate_review(
                text=all_text, topic=topic, citation=citation,
                length=length, tone=tone,
                inc_lim=inc_lim, inc_gaps=inc_gaps,
                inc_table=inc_table, inc_impl=inc_impl,
            )

            if not result:
                status.update(label="Generation failed.", state="error")
                st.error(st.session_state.last_error or "Generation failed. Check API key.")
                return

            status.update(label="Review generated!", state="complete")

        meta = dict(topic=topic, citation=citation, length=length, tone=tone,
                    inc_lim=inc_lim, inc_gaps=inc_gaps,
                    inc_table=inc_table, inc_impl=inc_impl,
                    file_count=len(uploaded))
        title = topic.strip() or f"Literature Review {now_str()}"
        save_review(title, result, [f.name for f in uploaded], meta)
        st.success(f"Saved — {word_count(result):,} words generated.")
        st.rerun()

    divider()

    # Step 4: Export
    page_head("Step 4 — Review &amp; export")
    if not st.session_state.cur_review:
        st.info("Generate a review above to unlock export.")
        return

    text  = st.session_state.cur_review
    meta  = st.session_state.cur_meta or {}
    title = meta.get("topic", "").strip() or "Literature Review"
    chips = meta_chips_html(meta)

    st.markdown(
        "<div class='s-card' style='display:flex;align-items:center;"
        "justify-content:space-between;flex-wrap:wrap;gap:8px;margin-bottom:14px;'>"
        "<div><strong>" + title + "</strong>"
        "<span class='chip' style='margin-left:8px;'>" + str(word_count(text)) + " words</span>"
        + chips + "</div></div>",
        unsafe_allow_html=True,
    )

    with st.expander("📄 Read full review", expanded=True):
        st.markdown(text)

    fname = slugify(title) + "_" + date_slug()
    d1, d2 = st.columns(2)
    with d1:
        st.download_button("📥 Download TXT", data=export_txt(text),
                           file_name=fname + ".txt", mime="text/plain",
                           use_container_width=True)
    with d2:
        try:
            st.download_button("📄 Download DOCX", data=export_docx(text, title),
                               file_name=fname + ".docx",
                               mime="application/vnd.openxmlformats-officedocument"
                                    ".wordprocessingml.document",
                               use_container_width=True)
        except Exception as ex:
            st.warning(f"DOCX export issue: {ex}")

    n1, n2 = st.columns(2)
    with n1:
        if st.button("➕ Start new review", use_container_width=True):
            st.session_state.cur_review = None
            st.session_state.cur_meta   = {}
            st.rerun()
    with n2:
        if st.button("📊 Go to Dashboard", use_container_width=True):
            nav_to("Dashboard")

    st.markdown("<div class='page-footer'>Scholara drafts only. "
                "Verify all citations before submitting.</div>", unsafe_allow_html=True)


# ── 11. DASHBOARD ────────────────────────────────────────────
def _parse_dt(s: str) -> datetime:
    try:    return datetime.strptime(s, "%Y-%m-%d %H:%M")
    except: return datetime.min

def dashboard_page() -> None:
    page_head("Dashboard", "All reviews — loaded from cloud (Airtable)")

    # Load from Airtable
    user_id = get_user_id()
    records = airtable_get_records("reviews", user_id)
    cloud_reviews = []
    for rec in records:
        fields = rec.get("fields", {})
        try:
            papers_list = json.loads(fields.get("papers", "[]"))
        except:
            papers_list = []
        try:
            meta_dict = json.loads(fields.get("meta", "{}"))
        except:
            meta_dict = {}
        cloud_reviews.append({
            "title": fields.get("title", "Untitled"),
            "date": fields.get("created_at", now_str()),
            "text": fields.get("review_text", ""),
            "papers": papers_list,
            "meta": meta_dict,
        })

    st.session_state.reviews = cloud_reviews
    reviews = st.session_state.reviews

    if not reviews:
        st.markdown(
            "<div class='s-card s-card-dash'>"
            "<div style='font-size:2.5rem;margin-bottom:10px;'>📂</div>"
            "<strong style='color:var(--slate-600);'>No reviews yet</strong><br/>"
            "<span style='font-size:.86rem;color:var(--slate-400);'>"
            "Use Demo Wizard to generate your first review.</span></div>",
            unsafe_allow_html=True,
        )
        if st.button("Go to Demo Wizard", type="primary"):
            nav_to("Demo Wizard")
        return

    total_papers = sum(len(r.get("papers", [])) for r in reviews)
    total_words  = sum(word_count(r.get("text", "")) for r in reviews)
    kpi_row([
        ("Total reviews",          str(len(reviews)),    "cloud-backed"),
        ("Total papers processed", str(total_papers),    "across all reviews"),
        ("Total output",           f"{total_words:,} w", "estimated word count"),
    ])

    divider()

    s1, s2, s3 = st.columns([2.2, 1.2, 0.7])
    with s1:
        st.session_state.search = st.text_input(
            "Search", value=st.session_state.search,
            placeholder="Title, topic, paper name…", label_visibility="collapsed")
    with s2:
        st.session_state.sort_by = st.selectbox(
            "Sort", SORT_OPTIONS,
            index=SORT_OPTIONS.index(st.session_state.sort_by)
                if st.session_state.sort_by in SORT_OPTIONS else 0,
            label_visibility="collapsed")
    with s3:
        if st.button("Clear", use_container_width=True):
            st.session_state.search  = ""
            st.session_state.sort_by = "Newest first"
            st.rerun()

    q = st.session_state.search.strip().lower()
    items = [r for r in reviews if not q or q in (
        r.get("title","") + r.get("text","") + " ".join(r.get("papers",[]))
    ).lower()]

    sb = st.session_state.sort_by
    rev = sb in ("Newest first", "Title Z-A", "Most papers")
    key_fn = {
        "Newest first": lambda r: _parse_dt(r.get("date","")),
        "Oldest first": lambda r: _parse_dt(r.get("date","")),
        "Title A-Z":    lambda r: r.get("title","").lower(),
        "Title Z-A":    lambda r: r.get("title","").lower(),
        "Most papers":  lambda r: len(r.get("papers",[])),
    }.get(sb, lambda r: _parse_dt(r.get("date","")))
    items = sorted(items, key=key_fn, reverse=rev)

    st.caption(f"Showing {len(items)} of {len(reviews)} review(s).")

    for i, r in enumerate(items):
        text    = r.get("text", "")
        title   = r.get("title", "Untitled")
        date    = r.get("date", "—")
        papers  = r.get("papers", [])
        preview = text[:PREVIEW_CHARS].rstrip() + ("…" if len(text) > PREVIEW_CHARS else "")

        st.markdown(
            "<div class='rev-card'>"
            "<div class='rev-title'>" + title + "</div>"
            "<div class='rev-meta'>" + date + " · " + str(len(papers)) + " paper(s) · "
            + str(word_count(text)) + " words</div>"
            + meta_chips_html(r.get("meta", {})) +
            "<div class='rev-preview'>" + preview + "</div>"
            "</div>",
            unsafe_allow_html=True,
        )

        rc1, rc2, rc3, rc4 = st.columns([1, 1, 1, 0.55])
        with rc1:
            if st.button("Open", key=f"open_{i}", use_container_width=True):
                st.session_state.cur_review = text
                st.session_state.cur_meta   = r.get("meta", {})
                nav_to("Demo Wizard")
        with rc2:
            st.download_button("↓ TXT", data=export_txt(text),
                               file_name=slugify(title) + ".txt", mime="text/plain",
                               key=f"txt_{i}", use_container_width=True)
        with rc3:
            try:
                st.download_button("↓ DOCX", data=export_docx(text, title),
                                   file_name=slugify(title) + ".docx",
                                   mime="application/vnd.openxmlformats-officedocument"
                                        ".wordprocessingml.document",
                                   key=f"docx_{i}", use_container_width=True)
            except Exception:
                st.caption("DOCX N/A")
        with rc4:
            if st.button("✕", key=f"del_{i}", help="Remove this review"):
                st.session_state.reviews = [
                    x for x in st.session_state.reviews
                    if not (x.get("title") == title and x.get("date") == date)
                ]
                st.rerun()

    divider()
    b1, b2 = st.columns(2)
    with b1:
        if st.button("➕ New Review", type="primary", use_container_width=True):
            st.session_state.cur_review = None
            st.session_state.cur_meta   = {}
            nav_to("Demo Wizard")
    with b2:
        if st.button("🗑️ Clear All", use_container_width=True):
            st.session_state.reviews    = []
            st.session_state.cur_review = None
            st.rerun()

    st.markdown("<div class='page-footer'>Cloud-backed storage — reviews persist across sessions.</div>",
                unsafe_allow_html=True)


# ── 12. HOW IT WORKS ─────────────────────────────────────────
def how_it_works_page() -> None:
    page_head("How It Works", "Four steps from raw PDFs to an exportable literature review.")
    STEPS = [
        ("01","📤","Upload your papers",
         f"Batch-upload up to {MAX_FILES} PDFs. Scholara validates sizes and duplicates automatically.",
         "Tip: Text-based PDFs extract cleanly. Run OCR on scanned files first."),
        ("02","⚙️","Configure your output",
         "Choose citation style, review depth, tone, and optional sections (gaps, limitations, methods table).",
         "Tip: A specific research topic focuses the synthesis significantly."),
        ("03","🧠","Gemini reads and synthesises",
         "Text is extracted and sent to Gemini. It clusters themes, identifies contradictions, "
         "maps gaps, and writes a grounded review — citing only your papers.",
         "Note: Generation takes 20-90 seconds depending on corpus size."),
        ("04","📝","Download and refine",
         "Download as .txt or .docx (headings pre-applied). Every review is saved to Dashboard.",
         "Reminder: Verify all citations against original papers before submitting."),
    ]
    for num, icon, title, body, note in STEPS:
        ca, cb = st.columns([0.17, 1], gap="medium")
        with ca:
            st.markdown(
                "<div style='text-align:center;padding-top:8px;'>"
                "<div style='font-family:\"DM Serif Display\",Georgia,serif;"
                "font-size:2.4rem;color:var(--slate-200);line-height:1;'>" + num + "</div>"
                "<div style='font-size:1.7rem;margin-top:3px;'>" + icon + "</div>"
                "</div>", unsafe_allow_html=True)
        with cb:
            st.markdown(
                "<div class='s-card s-card-accent'>"
                "<strong style='font-size:.98rem;color:var(--ink);'>" + title + "</strong>"
                "<p style='font-size:.86rem;color:var(--slate-600);line-height:1.65;"
                "margin:8px 0 8px;'>" + body + "</p>"
                "<span style='font-size:.78rem;color:var(--slate-400);font-style:italic;'>"
                + note + "</span></div>", unsafe_allow_html=True)
        st.write("")

    divider()
    _, cta, _ = st.columns([1, 1.2, 1])
    with cta:
        if st.button("Try the Demo Wizard", type="primary", use_container_width=True):
            nav_to("Demo Wizard")
    st.markdown("<div class='page-footer'>© 2026 Scholara</div>", unsafe_allow_html=True)


# ── 13. FEATURES ─────────────────────────────────────────────
def features_page() -> None:
    page_head("Features", "Everything built in — no extensions required.")
    FEATS = [
        ("📤","Multi-PDF upload", f"Up to {MAX_FILES} files, {MAX_TOTAL_MB} MB total."),
        ("🧠","Gemini-powered synthesis", "Theme clustering, contradiction detection, gap mapping."),
        ("🎓","5 citation styles","APA 7, IEEE, MLA 9, Chicago 17, Vancouver."),
        ("⚙️","Depth & tone", "Short to Detailed. Academic, Concise, Critical, Neutral."),
        ("🔍","Gap analysis", "Optional section mapping underexplored areas."),
        ("📊","Methods table", "Plain-text comparison table of methodologies."),
        ("📁","Cloud dashboard","Searchable, sortable review history, backed by Airtable."),
        ("📝","Export-ready", ".txt and .docx with heading structure applied."),
        ("♿","Accessible", "Keyboard nav, focus outlines, skip link, semantic HTML."),
    ]
    ca, cb, cc = st.columns(3, gap="medium")
    for col, (icon, title, desc) in zip([ca,cb,cc,ca,cb,cc,ca,cb,cc], FEATS):
        with col:
            st.markdown(
                "<div class='s-card' style='min-height:150px;'>"
                "<div style='font-size:1.5rem;margin-bottom:8px;'>" + icon + "</div>"
                "<strong style='font-size:.90rem;color:var(--ink);'>" + title + "</strong>"
                "<p style='font-size:.80rem;color:var(--slate-600);margin-top:5px;"
                "line-height:1.5;'>" + desc + "</p></div>",
                unsafe_allow_html=True)
    st.markdown("<div class='page-footer'>© 2026 Scholara</div>", unsafe_allow_html=True)


# ── 14. USE CASES ────────────────────────────────────────────
def use_cases_page() -> None:
    page_head("Use Cases", "Who uses Scholara and how.")
    CASES = [
        ("🎓","MS / PhD Students","var(--indigo-dim)","var(--indigo)",
         "Ayesha, a 2nd-year PhD at NUST, uploaded 30 papers on federated learning. "
         "In 5 minutes she had a 1,400-word structured review for her thesis Chapter 2.",
         ["Thesis Chapter 2","Gap analysis before proposals","Fast orientation in new sub-fields","Systematic reviews for seminars"]),
        ("🔬","Research Assistants","#ECFDF5","var(--success)",
         "Hamza, an RA at LUMS, was tasked with a bibliography for a funded HCI project. "
         "He uploaded 45 PDFs and got a themed synthesis in under 10 minutes.",
         ["Rapid literature mapping for PIs","Annotated bibliographies","Weekly research digests","Updating reviews with new papers"]),
        ("👨‍🏫","Assistant Professors","var(--amber-lt)","var(--amber)",
         "Dr. Usman uses Scholara for grant backgrounds and student paper supervision. "
         "It removes 6-8 hours from every submission cycle.",
         ["Grant proposal backgrounds","Journal submission related work","Student paper supervision","Staying current across sub-fields"]),
    ]
    for icon, audience, bg, border, scenario, uses in CASES:
        cl, cr = st.columns([1.1, 1], gap="large")
        with cl:
            st.markdown(
                "<div class='s-card' style='background:" + bg + ";border-left:4px solid "
                + border + ";padding-left:18px;'>"
                "<div style='font-size:1.8rem;margin-bottom:7px;'>" + icon + "</div>"
                "<strong style='font-size:.98rem;color:var(--ink);'>" + audience + "</strong>"
                "<p style='font-size:.85rem;color:var(--slate-600);margin-top:7px;"
                "line-height:1.6;'>" + scenario + "</p></div>",
                unsafe_allow_html=True)
        with cr:
            rows = "".join(
                "<li style='padding:5px 0;border-bottom:1px solid var(--slate-100);"
                "font-size:.84rem;color:var(--slate-600);'>"
                "<span style='color:var(--indigo);font-weight:700;'>→</span> " + u + "</li>"
                for u in uses)
            st.markdown(
                "<div class='s-card' style='height:100%;'>"
                "<p style='font-size:.73rem;font-weight:700;letter-spacing:.08em;"
                "text-transform:uppercase;color:var(--slate-400);margin-bottom:7px;'>"
                "Common uses</p>"
                "<ul style='list-style:none;padding:0;margin:0;'>" + rows + "</ul>"
                "</div>", unsafe_allow_html=True)
        st.write("")

    divider()
    st.markdown(
        "<div class='s-card s-card-dark' style='text-align:center;padding:32px;'>"
        "<p style='font-family:\"DM Serif Display\",Georgia,serif;font-size:1.4rem;"
        "color:var(--white);margin-bottom:8px;'>Your next review starts here.</p>"
        "<p style='color:#94A3B8;font-size:.88rem;margin-bottom:18px;'>"
        "No account required — upload and generate in minutes.</p></div>",
        unsafe_allow_html=True)
    _, cta, _ = st.columns([1,1,1])
    with cta:
        if st.button("Try the Demo Wizard", type="primary", use_container_width=True):
            nav_to("Demo Wizard")
    st.markdown("<div class='page-footer'>© 2026 Scholara</div>", unsafe_allow_html=True)


# ── 15. PRICING ──────────────────────────────────────────────
def pricing_page() -> None:
    page_head("Simple, transparent pricing", "Start free — no credit card needed.")
    PLANS = [
        ("free",  "Free",   "$0",  "forever",   False,
         [("✓","3 reviews per month"),("✓","Up to 10 PDFs/review"),
          ("✓","APA 7 only"),("✓","TXT export"),("✗","DOCX export"),("✗","All citation styles")]),
        ("scholar","Scholar","$9", "per month",  True,
         [("✓","Unlimited reviews"),("✓","Up to 50 PDFs/review"),
          ("✓","All 5 citation styles"),("✓","TXT + DOCX export"),
          ("✓","Review dashboard"),("✓","Research gap section")]),
        ("lab",   "Lab",   "$29", "per month",  False,
         [("✓","Everything in Scholar"),("✓","Up to 5 team members"),
          ("✓","Lab admin dashboard"),("✓","Priority processing"),
          ("✓","University invoice"),("✓","Dedicated support")]),
    ]
    cols = st.columns(3, gap="medium")
    for col, (key, name, price, period, featured, feats) in zip(cols, PLANS):
        with col:
            feat_style = "price-card featured" if featured else "price-card"
            badge = ("<div style='background:var(--indigo);color:#fff;font-size:.70rem;"
                     "font-weight:700;letter-spacing:.08em;text-transform:uppercase;"
                     "border-radius:999px;padding:3px 12px;display:inline-block;"
                     "margin-bottom:10px;'>Most popular</div>" if featured else "")
            rows = "".join(
                "<div class='price-row'>"
                "<span style='color:" + ("var(--success)" if t=="✓" else "var(--slate-200)") + ";"
                "font-weight:700;min-width:14px;'>" + t + "</span>"
                "<span style='color:" + ("var(--slate-600)" if t=="✓" else "var(--slate-200)") + ";"
                "font-size:.82rem;'>" + lbl + "</span></div>"
                for t, lbl in feats)
            st.markdown(
                "<div class='" + feat_style + "'>" + badge +
                "<div class='price-name'>" + name + "</div>"
                "<div class='price-val'>" + price + "</div>"
                "<div class='price-per'>" + period + "</div>"
                + rows + "</div>",
                unsafe_allow_html=True)
            st.write("")
            if featured:
                st.button("Start 7-day trial", key="plan_" + key,
                          type="primary", use_container_width=True)
            else:
                st.button(("Start free" if key=="free" else "Contact sales"),
                          key="plan_" + key, use_container_width=True)

    divider()
    d1, d2 = st.columns(2)
    with d1:
        st.markdown(
            "<div class='s-card s-card-amber'><strong>🌍 LMIC pricing</strong><br/>"
            "<span style='font-size:.85rem;color:var(--slate-600);'>"
            "Researchers in Pakistan and South Asia are eligible for a "
            "<strong>60% discount</strong> on all paid plans.</span></div>",
            unsafe_allow_html=True)
    with d2:
        st.markdown(
            "<div class='s-card s-card-accent'><strong>🏫 Annual billing</strong><br/>"
            "<span style='font-size:.85rem;color:var(--slate-600);'>"
            "Pay yearly and save <strong>20%</strong>. University POs accepted for Lab tier."
            "</span></div>", unsafe_allow_html=True)

    divider()
    page_head("FAQ")
    FAQ = [
        ("Does my uploaded PDF data get stored?",
         "No. Papers are processed in-session only. The session resets when you close the tab."),
        ("Can I use the output in my thesis?",
         "Yes — with care. Verify all citations against originals and revise before submission."),
        ("What if my PDF is scanned?",
         "Run OCR first (Adobe Acrobat, Smallpdf, or open-source OCRmyPDF), then upload."),
        ("Which AI model powers Scholara?",
         "Gemini 1.5 Flash, chosen for academic reading and synthesis strength."),
    ]
    for q, a in FAQ:
        with st.expander(q):
            st.markdown(f"<p style='font-size:.87rem;color:var(--slate-600);line-height:1.6;'>{a}</p>",
                        unsafe_allow_html=True)
    st.markdown("<div class='page-footer'>© 2026 Scholara · Pricing subject to change.</div>",
                unsafe_allow_html=True)


# ── 16. ABOUT ────────────────────────────────────────────────
def about_page() -> None:
    page_head("About Scholara")
    st.markdown(
        "<div class='s-card s-card-dark' style='display:flex;gap:18px;flex-wrap:wrap;'>"
        "<div style='font-size:2.8rem;'>👤</div><div>"
        "<strong style='color:#fff;font-size:.98rem;'>Soban — Founder</strong>"
        "<p style='color:#94A3B8;font-size:.86rem;margin-top:6px;line-height:1.65;max-width:540px;'>"
        "MS Artificial Intelligence candidate at NUST, Islamabad. "
        "Scholara grew from a personal frustration: spending 40+ hours on literature "
        "review for every thesis chapter while actual research got squeezed into "
        "whatever time was left.</p></div></div>",
        unsafe_allow_html=True)
    divider()
    c1, c2 = st.columns(2)
    with c1:
        page_head("Mission")
        st.markdown("<p style='font-size:.88rem;color:var(--slate-600);line-height:1.7;'>"
                    "Help researchers spend less time formatting and more time thinking. "
                    "Scholara recovers hours lost in manual literature review so you can "
                    "focus on ideas that actually advance your field.</p>", unsafe_allow_html=True)
    with c2:
        page_head("What we are not")
        st.markdown("<p style='font-size:.88rem;color:var(--slate-600);line-height:1.7;'>"
                    "A tool that writes your research for you. Scholara drafts; you decide. "
                    "Every output should be verified against originals and revised before "
                    "any academic submission.</p>", unsafe_allow_html=True)
    divider()
    page_head("Principles")
    PRINCIPLES = [
        ("🔒","Privacy by default","Reviews stored securely in your Airtable base."),
        ("📎","Source-grounded","Gemini synthesises only from your uploaded papers."),
        ("✏️","Draft, not final","All output is a starting point — review and revise."),
        ("♿","Accessible","Keyboard nav, focus outlines, skip link, semantic HTML."),
    ]
    p1,p2,p3,p4 = st.columns(4, gap="small")
    for col, (icon, title, body) in zip([p1,p2,p3,p4], PRINCIPLES):
        with col:
            st.markdown(
                "<div class='s-card' style='text-align:center;min-height:180px;'>"
                "<div style='font-size:1.7rem;margin-bottom:9px;'>" + icon + "</div>"
                "<strong style='font-size:.86rem;color:var(--ink);'>" + title + "</strong>"
                "<p style='font-size:.78rem;color:var(--slate-600);margin-top:5px;"
                "line-height:1.5;'>" + body + "</p></div>", unsafe_allow_html=True)
    divider()
    st.markdown(
        "<div class='s-card s-card-amber'><strong>Feedback</strong><br/>"
        "<span style='font-size:.85rem;color:var(--slate-600);'>"
        "Bug reports, feature requests, or institutional licensing: "
        "<strong>scholara@nust.edu.pk</strong> (placeholder).</span></div>",
        unsafe_allow_html=True)
    st.markdown("<div class='page-footer'>© 2026 Scholara · v" + APP_VERSION
                + " · Built at NUST, Islamabad</div>", unsafe_allow_html=True)


# ── 17. BLOG ─────────────────────────────────────────────────
def blog_page() -> None:
    page_head("Resources & Blog", "Practical guides for researchers — no fluff.")
    POSTS = [
        ("How to write a strong literature review for a thesis","8 min","Guide","chip",
         "Structure, thematic synthesis, gap identification, and the mistakes that get Chapter 2 sent back.","pub"),
        ("AI in academic writing: what researchers need to know","6 min","Analysis","chip chip-amber",
         "What AI tools can and cannot do — including where they hallucinate and how to catch it.","pub"),
        ("APA 7 vs IEEE: a practical guide","4 min","Reference","chip chip-green",
         "When to use which, key differences, and edge cases neither style guide covers clearly.","pub"),
        ("Finding real research gaps — not just 'more work is needed'","7 min","Guide","chip",
         "Contradiction mapping, methodological blind spots, and what actually impresses reviewers.","pub"),
        ("Scholara updates — July 2026","2 min","Changelog","chip chip-amber",
         "Redesigned UI, faster extraction, DOCX export with headings, LMIC pricing.","pub"),
        ("Reproducible literature review pipelines for research teams","Coming soon","Deep dive","chip",
         "Version control, shared corpora, and review assignment workflows.","upcoming"),
    ]
    for i, (title, read, tag, tag_cls, summary, status) in enumerate(POSTS, 1):
        num   = str(i).zfill(2) if status == "pub" else "—"
        opacity = "1" if status == "pub" else "0.60"
        st.markdown(
            "<div class='blog-card' style='opacity:" + opacity + ";'>"
            "<div class='blog-num'>" + num + "</div>"
            "<div style='flex:1;'>"
            "<div class='blog-title'>" + title + "</div>"
            "<div class='blog-meta' style='margin:4px 0 7px;'>"
            "<span class='" + tag_cls + "'>" + tag + "</span>&nbsp;·&nbsp;" + read + "</div>"
            "<p style='font-size:.82rem;color:var(--slate-600);line-height:1.55;margin:0;'>"
            + summary + "</p></div></div>",
            unsafe_allow_html=True)
    st.markdown("<div class='page-footer'>© 2026 Scholara</div>", unsafe_allow_html=True)


# ── 18. ROUTER ───────────────────────────────────────────────
PAGE_FN = {
    "Home":         home_page,
    "Demo Wizard":  demo_wizard_page,
    "Dashboard":    dashboard_page,
    "How It Works": how_it_works_page,
    "Features":     features_page,
    "Use Cases":    use_cases_page,
    "Pricing":      pricing_page,
    "About":        about_page,
    "Blog":         blog_page,
}

def load_reviews_from_cloud():
    """Pre-populate session state from Airtable on app start."""
    if not st.session_state.reviews:
        user_id = get_user_id()
        records = airtable_get_records("reviews", user_id)
        cloud_reviews = []
        for rec in records:
            fields = rec.get("fields", {})
            try:
                papers_list = json.loads(fields.get("papers", "[]"))
            except:
                papers_list = []
            try:
                meta_dict = json.loads(fields.get("meta", "{}"))
            except:
                meta_dict = {}
            cloud_reviews.append({
                "title": fields.get("title", "Untitled"),
                "date": fields.get("created_at", now_str()),
                "text": fields.get("review_text", ""),
                "papers": papers_list,
                "meta": meta_dict,
            })
        if cloud_reviews:
            st.session_state.reviews = cloud_reviews

def main() -> None:
    load_reviews_from_cloud()
    if st.session_state.page not in PAGE_FN:
        st.session_state.page = "Home"
    _sidebar()
    PAGE_FN[st.session_state.page]()

main()